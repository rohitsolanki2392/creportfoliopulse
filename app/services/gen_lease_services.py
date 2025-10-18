
import logging
import os
import json

from typing import Dict, Any
import uuid
from docx import Document
from fastapi import HTTPException
from sqlalchemy.orm import Session
from app.models.models import StandaloneFile, User
from datetime import datetime

from app.services.prompts import lease_abstract_prompt
import json
from typing import Dict, Any
from google import genai
import os
from sqlalchemy.orm import Session
from fastapi import HTTPException
from app.models.models import StandaloneFile
client = genai.Client()
from app.utils.process_file import extract_text_from_file, save_to_temp

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s")

with open("app/services/templates/lease_template.txt", "r", encoding="utf-8") as f:
    LEASE_TEMPLATE = f.read()

def save_lease_file(content: str, company_id: str, category: str, file_id: str) -> str:
    try:
        file_path = f"uploads/{company_id}/{category}/{file_id}.txt"
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        return file_path
    except Exception as e:
        logger.error(f"Error saving lease file {file_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save lease file: {str(e)}")





def generate_lease_text(metadata: Dict[str, Any]) -> str:
    lease_text = LEASE_TEMPLATE
    try:
        prompt = f"""
        You are an expert in lease document automation.
        Analyze the following lease template and metadata, and determine which metadata keys should replace which placeholders.
        Use your best judgment based on context (e.g., names, dates, addresses, terms).

        Template:
        {LEASE_TEMPLATE}

        Metadata:
        {metadata}

        Return only a valid JSON mapping of placeholder to value. Example:
        {{
          "[TENANT_NAME]": "John Smith",
          "[LANDLORD_NAME]": "ABC Properties",
          "[PROPERTY_ADDRESS]": "123 Main St, New York"
        }}
        """

        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=prompt
        )

        raw_text = response.text.strip()

        # ✅ Remove Markdown code fences if present
        if raw_text.startswith("```"):
            raw_text = raw_text.strip("`")
            raw_text = raw_text.replace("json", "", 1).strip()

        # ✅ Safely parse JSON
        try:
            replacements = json.loads(raw_text)
        except Exception as json_error:
            print("Raw response from Gemini:", raw_text)
            raise ValueError(f"Gemini returned invalid JSON: {json_error}")

        for placeholder, value in replacements.items():
            lease_text = lease_text.replace(placeholder, str(value or "N/A"))

        return lease_text

    except Exception as e:
        return f"Error processing lease template with metadata: {e}"



async def process_lease_abstract(file, current_user, category: str, db: Session):

    try:
 
        temp_path = await save_to_temp(file, current_user.id, current_user, category)
        extracted_text = extract_text_from_file(temp_path)
        if not extracted_text:
            raise HTTPException(400, "No text extracted from file")

        prompt = lease_abstract_prompt + "\n\nLEASE TEXT:\n" + extracted_text
        response = client.models.generate_content(model="gemini-2.0-flash", contents=prompt)
        abstract_text = response.text.strip()


        doc = Document()
        doc.add_heading("Commercial Office Lease Abstract", level=1)
        for line in abstract_text.split("\n"):
            if line.strip():
                if line.startswith(("I.", "II.", "III.")):
                    doc.add_heading(line.strip(), level=2)
                else:
                    doc.add_paragraph(line.strip())

        base_dir = os.path.join("uploads", str(current_user.company_id), category)
        os.makedirs(base_dir, exist_ok=True)
        output_name = f"lease_abstract_{uuid.uuid4().hex}.docx"
        output_path = os.path.join(base_dir, output_name)
        doc.save(output_path)


        file_size = os.path.getsize(output_path)
        new_file = StandaloneFile(
            file_id=str(uuid.uuid4()),
            original_file_name=file.filename,
            user_id=current_user.id,
            company_id=current_user.company_id,
            category=category,
            uploaded_at=datetime.utcnow(),
            gcs_path=output_path, 
            file_size=str(file_size),
        )
        db.add(new_file)
        db.commit()
        db.refresh(new_file)

        return output_path

    except Exception as e:
        logger.error(f"Error processing lease abstract: {e}")
        raise HTTPException(500, f"Failed to process file: {e}")






async def list_category_files_service(user_id: str, category: str, db: Session):
    try:
        files = db.query(StandaloneFile).filter(
            StandaloneFile.user_id == user_id,
            StandaloneFile.category == category,
        ).all()

        result = []
        for file in files:
            # Convert local path to web URL
            normalized_path = file.gcs_path.replace("\\", "/")
            full_url = f"/{normalized_path}"

            result.append({
                "file_id": file.file_id,
                "original_file_name": file.original_file_name,
                "user_id": file.user_id,
                "uploaded_at": file.uploaded_at.isoformat(),
                "category": file.category,
                "file_url": full_url  # <-- changed from file_path to file_url
            })

        return {
            "category": category,
            "files": result,
            "total_files": len(result),
        }

    except Exception as e:
        return {"error": str(e)}



async def delete_file_service(file_id: str, user_id: str, db: Session):
    try:
        file_record = db.query(StandaloneFile).filter(
            StandaloneFile.file_id == file_id,
            StandaloneFile.user_id == user_id,
        ).first()

        if not file_record:
            raise HTTPException(status_code=404, detail="File not found")

        # Delete local file if exists
        if os.path.exists(file_record.gcs_path):
            os.remove(file_record.gcs_path)

        db.delete(file_record)
        db.commit()

        return {"message": f"File '{file_record.original_file_name}' deleted successfully"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")
