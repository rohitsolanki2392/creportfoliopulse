import os
import uuid
import logging
from datetime import datetime
from docx import Document
from fastapi import HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from app.models.models import StandaloneFile
from app.utils.process_file import extract_text_from_file, save_to_temp
from app.services.prompts import lease_abstract_prompt
from google import genai
import asyncio
from sqlalchemy.future import select

client = genai.Client()
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

async def save_lease_file(content: str, company_id: str, category: str, file_id: str) -> str:
    try:
        file_path = f"uploads/{company_id}/{category}/{file_id}.txt"
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        await asyncio.to_thread(lambda: open(file_path, "w", encoding="utf-8").write(content))
        return file_path
    except Exception as e:
        logger.error(f"Error saving lease file {file_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save lease file: {str(e)}")


async def process_lease_abstract(file, current_user, category: str, db: AsyncSession):
    try:

        temp_path = await save_to_temp(file, current_user.id, current_user, category)

        extracted_text = await extract_text_from_file(temp_path)
        if not extracted_text:
            raise HTTPException(400, "No text extracted from file")

        prompt = lease_abstract_prompt + "\n\nLEASE TEXT:\n" + extracted_text
        response = await asyncio.to_thread(lambda: client.models.generate_content(
            model="gemini-2.0-flash",
            contents=prompt
        ))
        abstract_text = response.text.strip()


        doc = Document()
        doc.add_heading("Commercial Office Lease Abstract", level=1)
        for line in abstract_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith(("I.", "II.", "III.")):
                doc.add_heading(line, level=2)
            else:
                doc.add_paragraph(line)


        base_dir = os.path.join("uploads", str(current_user.company_id), category)
        os.makedirs(base_dir, exist_ok=True)
        output_name = f"lease_abstract_{uuid.uuid4().hex}.docx"
        output_path = os.path.join(base_dir, output_name)
        await asyncio.to_thread(doc.save, output_path)


        file_size = os.path.getsize(output_path)
        new_file = StandaloneFile(
            file_id=str(uuid.uuid4()),
            original_file_name=file.filename,
            user_id=current_user.id,
            company_id=current_user.company_id,
            category=category,
            uploaded_at=datetime.utcnow(),
            gcs_path=output_path,
            file_size=str(file_size),
        )
        db.add(new_file)
        await db.commit()
        await db.refresh(new_file)

        return output_path

    except Exception as e:
        logger.error(f"Error processing lease abstract: {e}")
        raise HTTPException(500, f"Failed to process file: {e}")


async def list_category_files_service(user_id: str, category: str, db: AsyncSession):
    try:
        result = []

        files_result = await db.execute(
            select(StandaloneFile).where(
                StandaloneFile.user_id == user_id,
                StandaloneFile.category == category
            )
        )
        files = files_result.scalars().all()

        for file in files:

            result.append({
                "file_id": file.file_id,
                "original_file_name": file.original_file_name,
                "user_id": file.user_id,
                "uploaded_at": file.uploaded_at.isoformat(),
                "category": file.category,
                "file_url": None
            })

        return {
            "category": category,
            "files": result,
            "total_files": len(result)
        }

    except Exception as e:
        return {"error": str(e)}



async def delete_file_service(file_id: str, user_id: str, db: AsyncSession):
    try:

        file_record = await db.get(StandaloneFile, file_id)
        if not file_record or file_record.user_id != user_id:
            raise HTTPException(status_code=404, detail="File not found")



        await db.delete(file_record)
        await db.commit()

        return {"message": f"File '{file_record.original_file_name}' deleted successfully"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

